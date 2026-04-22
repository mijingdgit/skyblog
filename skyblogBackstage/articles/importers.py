from __future__ import annotations

from io import BytesIO
import json
from pathlib import Path
import re
from urllib.parse import unquote, urlparse
import uuid
import zipfile
from xml.etree import ElementTree

from django import forms
from django.core.files.base import ContentFile
from django.core.files.storage import default_storage
from django.utils.text import slugify

try:
    from docx import Document
except ImportError:  # pragma: no cover
    Document = None


SUPPORTED_IMPORT_SUFFIXES = {".md", ".markdown", ".xmind", ".docx", ".doc"}
MARKDOWN_IMAGE_PATTERN = re.compile(r"!\[([^\]]*)\]\(([\s\S]*?)\)")
HTML_IMAGE_PATTERN = re.compile(r'(<img\b[^>]*?\bsrc=["\'])([^"\']+)(["\'][^>]*>)', re.IGNORECASE)
XMIND_SOURCE_MARKER = "<!-- source:xmind -->"


def extract_content_from_upload(uploaded) -> str:
    suffix = Path(uploaded.name).suffix.lower()
    payload = uploaded.read()

    if suffix in {".md", ".markdown"}:
        content = _decode_text_payload(payload)
        return normalize_markdown_assets(content, strict=True)
    if suffix == ".xmind":
        return _extract_xmind_content(payload)
    if suffix == ".docx":
        return _extract_docx_content(payload)
    if suffix == ".doc":
        return _extract_doc_content(payload)

    raise forms.ValidationError("仅支持导入 .md、.markdown、.xmind、.docx 或 .doc 文件。")


def normalize_markdown_assets(content: str, strict: bool = False) -> str:
    rewritten = _rewrite_markdown_images(content, strict=strict)
    rewritten = _rewrite_html_images(rewritten, strict=strict)
    return rewritten


def infer_title(content: str, filename: str) -> str:
    for line in content.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped == XMIND_SOURCE_MARKER:
            continue
        if stripped.startswith("#"):
            return stripped.lstrip("#").strip() or Path(filename).stem
        if len(stripped) <= 80:
            return stripped
        break

    return Path(filename).stem


def infer_excerpt(content: str) -> str:
    lines: list[str] = []
    for line in content.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped == XMIND_SOURCE_MARKER:
            continue
        if stripped.startswith("#") or stripped.startswith("```") or stripped.startswith("!"):
            continue
        lines.append(stripped)

    excerpt = " ".join(lines)
    excerpt = _strip_markdown(excerpt)
    excerpt = re.sub(r"\s+", " ", excerpt).strip()
    return excerpt[:140] if excerpt else "由导入文件自动生成的文章内容。"


def _decode_text_payload(payload: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "gbk"):
        try:
            return payload.decode(encoding)
        except UnicodeDecodeError:
            continue

    raise forms.ValidationError("文件编码无法识别，请优先使用 UTF-8 或 GBK 编码后重试。")


def _extract_xmind_content(payload: bytes) -> str:
    try:
        with zipfile.ZipFile(BytesIO(payload)) as archive:
            if "content.json" in archive.namelist():
                data = json.loads(archive.read("content.json").decode("utf-8"))
                return f"{XMIND_SOURCE_MARKER}\n\n{_xmind_json_to_markdown(data)}"

            if "content.xml" in archive.namelist():
                xml_payload = archive.read("content.xml")
                return f"{XMIND_SOURCE_MARKER}\n\n{_xmind_xml_to_markdown(xml_payload)}"
    except zipfile.BadZipFile as exc:
        raise forms.ValidationError("XMind 文件无法识别，请确认上传的是有效的 .xmind 文件。") from exc

    raise forms.ValidationError("当前 XMind 文件缺少可读取的主题内容。")


def _xmind_json_to_markdown(data) -> str:
    sheets = data if isinstance(data, list) else [data]
    outlines: list[str] = []

    for sheet in sheets:
        root = sheet.get("rootTopic") or {}
        title = (root.get("title") or sheet.get("title") or "思维导图").strip()
        outlines.append(f"# {title}")
        outlines.extend(_walk_xmind_json_children(root, 0))

    return "\n".join(line for line in outlines if line).strip()


def _walk_xmind_json_children(topic: dict, depth: int) -> list[str]:
    lines: list[str] = []
    title = (topic.get("title") or "").strip()
    if depth > 0 and title:
        lines.append(f'{"  " * (depth - 1)}- {title}')

    children = topic.get("children") or {}
    attached = children.get("attached") or []
    if isinstance(attached, dict):
        attached = attached.get("topics") or []

    for child in attached:
        lines.extend(_walk_xmind_json_children(child, depth + 1))

    return lines


def _xmind_xml_to_markdown(payload: bytes) -> str:
    try:
        root = ElementTree.fromstring(payload)
    except ElementTree.ParseError as exc:
        raise forms.ValidationError("XMind XML 内容解析失败，请检查文件是否完整。") from exc

    lines: list[str] = []
    for sheet in root.findall(".//{*}sheet"):
        topic = sheet.find("{*}topic")
        if topic is None:
            continue
        title = _xml_topic_title(topic) or "思维导图"
        lines.append(f"# {title}")
        lines.extend(_walk_xmind_xml_children(topic, 0))

    if not lines:
        raise forms.ValidationError("XMind 文件中没有可导入的主题内容。")

    return "\n".join(lines).strip()


def _walk_xmind_xml_children(topic, depth: int) -> list[str]:
    lines: list[str] = []
    if depth > 0:
        title = _xml_topic_title(topic)
        if title:
            lines.append(f'{"  " * (depth - 1)}- {title}')

    for child in topic.findall("./{*}children/{*}topics/{*}topic"):
        lines.extend(_walk_xmind_xml_children(child, depth + 1))

    return lines


def _xml_topic_title(topic) -> str:
    title_node = topic.find("{*}title")
    if title_node is None or title_node.text is None:
        return ""
    return title_node.text.strip()


def _extract_docx_content(payload: bytes) -> str:
    if Document is None:
        raise forms.ValidationError("当前环境缺少 python-docx 依赖，无法读取 .docx 文件，请先安装依赖后重试。")

    document = Document(BytesIO(payload))
    blocks: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(text)

    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                blocks.append(row_text)

    content = "\n\n".join(blocks).strip()
    if not content:
        raise forms.ValidationError("Word 文档中没有可导入的正文内容。")

    return content


def _extract_doc_content(payload: bytes) -> str:
    utf16_text = payload.decode("utf-16le", errors="ignore")
    utf16_matches = re.findall(r"[\u4e00-\u9fffA-Za-z0-9，。！？；：《》（）【】、,.!?;:()\-_/ ]{4,}", utf16_text)
    ansi_matches = [match.decode("latin1").strip() for match in re.findall(rb"[\x20-\x7E]{4,}", payload)]

    merged: list[str] = []
    seen: set[str] = set()

    for chunk in utf16_matches + ansi_matches:
        normalized = re.sub(r"\s+", " ", chunk).strip()
        if len(normalized) < 4 or normalized in seen:
            continue
        seen.add(normalized)
        merged.append(normalized)

    content = "\n\n".join(merged).strip()
    if not content:
        raise forms.ValidationError("旧版 .doc 文件内容提取失败，请优先将文件另存为 .docx 后再导入。")

    return content


def _rewrite_markdown_images(content: str, strict: bool = False) -> str:
    missing_paths: list[str] = []

    def replace(match: re.Match[str]) -> str:
        alt_text = match.group(1)
        raw_path = match.group(2).strip()
        normalized_path = _materialize_local_asset(raw_path)
        if not normalized_path:
            if strict and _looks_like_local_file(raw_path):
                missing_paths.append(raw_path)
            return match.group(0)
        return f"![{alt_text}]({normalized_path})"

    rewritten = MARKDOWN_IMAGE_PATTERN.sub(replace, content)
    _raise_missing_asset_error(missing_paths)
    return rewritten


def _rewrite_html_images(content: str, strict: bool = False) -> str:
    missing_paths: list[str] = []

    def replace(match: re.Match[str]) -> str:
        prefix, raw_path, suffix = match.groups()
        normalized_path = _materialize_local_asset(raw_path)
        if not normalized_path:
            if strict and _looks_like_local_file(raw_path):
                missing_paths.append(raw_path)
            return match.group(0)
        return f"{prefix}{normalized_path}{suffix}"

    rewritten = HTML_IMAGE_PATTERN.sub(replace, content)
    _raise_missing_asset_error(missing_paths)
    return rewritten


def _materialize_local_asset(raw_path: str) -> str | None:
    path_text = raw_path.strip().strip("<>").strip().replace("\\\\", "\\")
    path_text = unquote(path_text)
    path_text = re.sub(r"\s+", " ", path_text)

    if path_text.startswith("file:///"):
        parsed = urlparse(path_text)
        path_text = unquote(parsed.path.lstrip("/"))

    if not _looks_like_local_file(path_text):
        return None

    source_path = Path(path_text)
    if not source_path.exists() or not source_path.is_file():
        return None

    suffix = source_path.suffix.lower()
    stem = slugify(source_path.stem) or "asset"
    target_name = f"{stem}-{uuid.uuid4().hex[:10]}{suffix}"
    storage_path = f"article-assets/{target_name}"

    with source_path.open("rb") as file_handle:
        saved_path = default_storage.save(storage_path, ContentFile(file_handle.read()))

    return default_storage.url(saved_path)


def _looks_like_local_file(path_text: str) -> bool:
    normalized = path_text.strip()
    if re.match(r"^[a-zA-Z]:[\\/]", normalized):
        return True
    if normalized.startswith("\\\\"):
        return True
    if normalized.startswith("/"):
        return True
    if normalized.startswith("file:///"):
        return True
    return False


def _raise_missing_asset_error(paths: list[str]) -> None:
    if not paths:
        return

    unique_paths = list(dict.fromkeys(paths))
    sample = "；".join(unique_paths[:3])
    raise forms.ValidationError(
        f"以下 Markdown 图片路径当前无法被 Django 访问：{sample}。请确认这些文件存在且当前 Python 进程可访问，或先把图片放到站点可读取的位置后再导入。"
    )


def _strip_markdown(value: str) -> str:
    cleaned = re.sub(r"!\[[^\]]*\]\([^)]*\)", "", value)
    cleaned = re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", cleaned)
    cleaned = re.sub(r"[#>*`~_-]", " ", cleaned)
    return cleaned
